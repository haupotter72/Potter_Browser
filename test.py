import sys
import json
import os
import socket
import json

from PyQt5.QtWidgets import (QApplication, QWidget, QVBoxLayout, QHBoxLayout,
                             QPushButton, QLabel, QLineEdit, QTabBar,
                             QFrame, QStackedLayout)

from PyQt5.QtGui import QIcon, QWindow, QImage
from PyQt5.QtCore import *
from PyQt5.QtWebEngineWidgets import *
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class AddressBar(QLineEdit):
    def __init__(self):
        super().__init__()

    def mousePressEvent(self, e):
        self.selectAll()


class App(QFrame):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Tho Thang Browser")

        self.CreateApp()
        self.setBaseSize(1280, 900)
        
       
           
  
        self.showMaximized()


    def CreateApp(self):
        self.layout = QVBoxLayout()
        self.layout.setSpacing(0)
        self.layout.setContentsMargins(0, 0, 0, 0)

        self.tabbar = QTabBar(movable=True, tabsClosable=True)
        self.tabbar.tabCloseRequested.connect(self.CloseTab)
        self.tabbar.tabBarClicked.connect(self.SwitchTab)
        
 


        self.tabbar.setCurrentIndex(0)
        self.tabbar.setDrawBase(False)

        self.tabCount = 0
        self.tabs = []


        self.Toolbar = QWidget()
        self.ToolbarLayout = QHBoxLayout()
        self.addressbar = AddressBar()

        self.AddTabButton = QPushButton(QIcon(os.path.join('images', 'ui-tab--plus.png')), "  New Tab", self)

        self.AddTabButton.clicked.connect(self.AddTab)

        self.addressbar.returnPressed.connect(self.BrowseTo)

        #Toolbar Buttons

        self.BackButton = QPushButton(QIcon(os.path.join('images', 'arrow-180.png')), "Back", self)
        self.BackButton.clicked.connect(self.GoBack)

        self.ForwardButton = QPushButton(QIcon(os.path.join('images', 'arrow-000.png')), " Forward", self)
        self.ForwardButton.clicked.connect(self.GoForward)

        self.ReloadButton = QPushButton(QIcon(os.path.join('images', 'arrow-circle-315.png')), " Reload", self)
        self.ReloadButton.clicked.connect(self.ReloadPage)

        # Toolbar
        self.ToolbarLayout.addWidget(self.AddTabButton)
        self.ToolbarLayout.addWidget(self.BackButton)
        self.ToolbarLayout.addWidget(self.ForwardButton)
        self.ToolbarLayout.addWidget(self.ReloadButton)
        self.Toolbar.setLayout(self.ToolbarLayout)
        self. ToolbarLayout.addWidget(self.addressbar)


        # Set Main View
        self.container = QWidget()
        self.container.layout = QStackedLayout()
        self.container.setLayout(self.container.layout)


        # Main View From Top level Elements
        self.layout.addWidget(self.tabbar)
        self.layout.addWidget(self.Toolbar)
        self.layout.addWidget(self.container)
        
        self.favorite_websites = []

        # Tạo một nút để thêm trang web vào danh sách yêu thích
        self.AddFavoriteButton = QPushButton(QIcon(os.path.join('images', 'heart.png')), " Add to Favorites", self)
        self.AddFavoriteButton.clicked.connect(self.AddFavoriteWebsite)
        self.ToolbarLayout.addWidget(self.AddFavoriteButton)
        
        # Tạo một nút để xóa trang web khỏi danh sách yêu thích
        self.RemoveFavoriteButton = QPushButton(QIcon(os.path.join('images', 'heart-broken.png')), " Remove from Favorites", self)
        self.RemoveFavoriteButton.clicked.connect(self.RemoveFavoriteWebsite)
        self.ToolbarLayout.addWidget(self.RemoveFavoriteButton)
        
        
       



        # Tạo một nút để hiển thị danh sách trang web yêu thích
        self.ShowFavoritesButton = QPushButton(" Show Favorites", self)
        self.ShowFavoritesButton.clicked.connect(self.ShowFavoriteWebsites)
        self.ToolbarLayout.addWidget(self.ShowFavoritesButton)
         # Tạo một tab mới cho danh sách trang web yêu thích
        self.favorites_tab = QWidget()
        self.favorites_tab.layout = QVBoxLayout()
        self.favorites_tab.layout.setContentsMargins(0, 0, 0, 0)
        self.favorites_tab.content = QWebEngineView()
        self.favorites_tab.layout.addWidget(self.favorites_tab.content)
        self.favorites_tab.setLayout(self.favorites_tab.layout)
        self.container.layout.addWidget(self.favorites_tab)
        self.favorites_tab.setVisible(False)  # Ẩn tab này ban đầu

        self.setLayout(self.layout)

        self.AddTab()

        self.show()

    def CloseTab(self, i):
        self.tabbar.removeTab(i)
    
    
    def AddFavoriteWebsite(self):
        # Lấy URL hiện tại từ thanh địa chỉ
        current_url = self.addressbar.text()

        # Kiểm tra xem URL đã tồn tại trong danh sách yêu thích chưa
        if current_url not in self.favorite_websites:
            self.favorite_websites.append(current_url)
            print(f"Added to favorites: {current_url}")
        else:
            print(f"This website is already in favorites: {current_url}")
    
    
    def ShowFavoriteWebsites(self):
     if self.favorite_websites:
        # Kiểm tra xem tệp DOCX đã tồn tại chưa
        try:
            doc = Document('favorite_websites.docx')
        except FileNotFoundError:
            # Nếu chưa tồn tại, tạo tệp mới
            doc = Document()
            doc.add_heading('Favorite Websites', 0)

        # Thêm danh sách trang web yêu thích vào tệp DOCX
        for website in self.favorite_websites:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = p.add_run(website)
            run.font.size = Pt(12)

        # Lưu tệp DOCX
        doc.save('favorite_websites.docx')

        # Hiển thị tab danh sách trang web yêu thích
        self.favorites_tab.setVisible(True)

        # Hiển thị danh sách yêu thích trong tab danh sách trang web yêu thích
        favorites_html = "<html><body><h1>Favorite Websites</h1><ul>"
        for website in self.favorite_websites:
            favorites_html += f"<li><a href=\"{website}\">{website}</a></li>"
        favorites_html += "</ul></body></html>"
        self.favorites_tab.content.setHtml(favorites_html)

        # Chuyển đến tab danh sách trang web yêu thích
        self.container.layout.setCurrentWidget(self.favorites_tab)
     else:
           self.favorites_tab.setVisible(False)
        
        
    def RemoveFavoriteWebsite(self):
    # Lấy URL hiện tại từ thanh địa chỉ
     current_url = self.addressbar.text()
     if current_url in self.favorite_websites:
        self.favorite_websites.remove(current_url)
        print(f"Removed from favorites: {current_url}")
        # Cập nhật tab danh sách trang web yêu thích
        self.ShowFavoriteWebsites()
     else:
        print(f"This website is not in favorites: {current_url}")






    def AddTab(self):
        i = self.tabCount

        self.tabs.append(QWidget())
        self.tabs[i].layout = QVBoxLayout()
        self.tabs[i].layout.setContentsMargins(0, 0, 0, 0)

        # To switch tab name through Index
        self.tabs[i].setObjectName("tab" + str(i))

        self.tabs[i].content = QWebEngineView()
        self.tabs[i].content.load(QUrl.fromUserInput("http://google.com"))

        # passing through Lambda the TabText to the Index
        self.tabs[i].content.titleChanged.connect(lambda: self.SetTabContent(i, "title"))
        self.tabs[i].content.iconChanged.connect(lambda: self.SetTabContent(i, "icon"))
        self.tabs[i].content.urlChanged.connect(lambda: self.SetTabContent(i, "url"))

        self.tabs[i].layout.addWidget(self.tabs[i].content)

        self.tabs[i].setLayout(self.tabs[i].layout)

        self.container.layout.addWidget(self.tabs[i])
        self.container.layout.setCurrentWidget(self.tabs[i])

        self.tabbar.addTab("Add New")
        self.tabbar.setTabData(i, {"object": "tab" + str(i), "initial": i})
        self.tabbar.setCurrentIndex(i)

        self.tabCount += 1

    def SwitchTab(self, i):

        if self.tabbar.tabData(i):

            tab_Data = self.tabbar.tabData(i)["object"]
            tab_Content = self.findChild(QWidget, tab_Data)
            self.container.layout.setCurrentWidget(tab_Content)

            new_url = tab_Content.content.url().toString()
            self.addressbar.setText(new_url)

    def BrowseTo(self):
        text = self.addressbar.text()
        print(text)

        i = self.tabbar.currentIndex()
        tab = self.tabbar.tabData(i)["object"]
        wv = self.findChild(QWidget, tab).content

        if "http" not in text:
            if "." not in text:
                url = "https://www.google.com/search?q=" + text #không nhập một tên miền cụ thể, thì biến url sẽ được gán giá trị
            else:
                url = "http://" + text

        else:
            url = text

        wv.load(QUrl.fromUserInput(url))
        
         # Thêm tab mới sau khi thực hiện tìm kiếm
        self.AddTab()

    def SetTabContent(self, i, type):
        tab_name = self.tabs[i].objectName()

        count = 0
        running = True

        current_tab = self.tabbar.tabData(self.tabbar.currentIndex())["object"]

        if current_tab == tab_name and type == "url":
            new_url = self.findChild(QWidget, tab_name).content.url().toString()
            self.addressbar.setText(new_url)
            return False

        while running:
            tab_data_name = self.tabbar.tabData(count)

            if count >= 10:
                running = False

            if tab_name == tab_data_name["object"]:
                if type == "title":
                    newTitle = self.findChild(QWidget, tab_name).content.title()
                    self.tabbar.setTabText(count, newTitle)
                elif type == "icon":
                    newIcon = self.findChild(QWidget, tab_name).content.icon()
                    self.tabbar.setTabIcon(count, newIcon)

                running = False
            else:
                count += 1

    def GoBack(self):
        activeIndex = self.tabbar.currentIndex()
        tab_name = self.tabbar.tabData(activeIndex)["object"]
        tab_content = self.findChild(QWidget, tab_name).content

        tab_content.back()

    def GoForward(self):
        activeIndex = self.tabbar.currentIndex()
        tab_name = self.tabbar.tabData(activeIndex)["object"]
        tab_content = self.findChild(QWidget, tab_name).content

        tab_content.forward()
        pass

    def ReloadPage(self):
        activeIndex = self.tabbar.currentIndex()
        tab_name = self.tabbar.tabData(activeIndex)["object"]
        tab_content = self.findChild(QWidget, tab_name).content

        tab_content.reload()
    
    def check_tcp_connection(self):
        try:
            # Thay đổi địa chỉ IP và port tương ứng
            sock = socket.create_connection(("example.com", 80))
            sock.close()
            return True
        except socket.error:
            return False    
        

# if __name__ == "__main__":
#     app = QApplication(sys.argv)
#     app.setApplicationName("Potter Browser")

#     window = App()

#     sys.exit(app.exec_())
app = QApplication(sys.argv)
app.setApplicationName("Potter Browser")



window = App()
#Kiểm tra kết nối TCP trước khi chạy ứng dụng
if window.check_tcp_connection():
    app.exec_()
else:
    print("Không thể kết nối TCP. Vui lòng kiểm tra kết nối mạng của bạn.")